from docx import Document
from docx.shared import Inches
import json
import os

def docx_fixer():
    os.chdir('data')
    with open('data.json', 'r', encoding='utf-8') as file:
        text = ''.join(file.readlines())[1:]
        data = json.loads(text)
    os.chdir('../docs_for_test')
    if file == None:
        doc = Document()
    else:
        doc = Document(f'{file}')
        
    for paragraph in doc.paragraphs:
        print(paragraph.text)
        
    os.chdir('../result')
    if file == None:
        pass
    else:
        doc.save(f'{file}')
    os.chdir('..')
   
## Проходим по всем абзацам в документе
#for paragraph in doc.paragraphs:
#    print(f'Текст: {paragraph.text}')
#    
#    # Проходим по всем runs (частям текста с одинаковым стилем) в абзаце
#    for run in paragraph.runs:
#        # Получаем параметры текста
#        font = run.font
#        print(f'  Шрифт: {font.name}')
#        print(f'  Размер: {font.size}')
#        print(f'  Жирный: {font.bold}')
#        print(f'  Курсив: {font.italic}')
#        print(f'  Подчеркнутый: {font.underline}')
#        
#    # Получаем отступы
#    print(f'  Отступ слева: {paragraph.paragraph_format.left_indent}')
#    print(f'  Отступ справа: {paragraph.paragraph_format.right_indent}')
#    print(f'  Отступ сверху: {paragraph.paragraph_format.space_before}')
#    print(f'  Отступ снизу: {paragraph.paragraph_format.space_after}')

## Получаем параметры отступов от края листа
#for section in doc.sections:
#    print(f'Отступ сверху: {section.top_margin}')
#    print(f'Отступ снизу: {section.bottom_margin}')
#    print(f'Отступ слева: {section.left_margin}')
#    print(f'Отступ справа: {section.right_margin}')
#
## Получаем параметры отступов от рамки до текста
#for paragraph in doc.paragraphs:
#    print(f'Текст: {paragraph.text}')
#    print(f'  Отступ слева: {paragraph.paragraph_format.left_indent}')
#    print(f'  Отступ справа: {paragraph.paragraph_format.right_indent}')
#    print(f'  Отступ сверху: {paragraph.paragraph_format.space_before}')
#    print(f'  Отступ снизу: {paragraph.paragraph_format.space_after}')

## Изменяем параметры отступов от края листа
#for section in doc.sections:
#    section.top_margin = 100000  # 100000 EMU (1 EMU = 1/1000000 дюйма)
#    section.bottom_margin = 100000
#    section.left_margin = 100000
#    section.right_margin = 100000
#
## Изменяем параметры отступов от рамки до текста
#for paragraph in doc.paragraphs:
#    paragraph.paragraph_format.left_indent = 20000  # 20000 EMU
#    paragraph.paragraph_format.right_indent = 20000
#    paragraph.paragraph_format.space_before = 1000
#    paragraph.paragraph_format.space_after = 1000